"""
Document ingestion pipeline for Arkive AI.

Supported formats: PDF, DOCX, TXT

Features:
- PyMuPDF text extraction with OCR fallback (Tesseract) for PDFs
- python-docx extraction for Word documents
- Plain text extraction for .txt files
- Password-protected PDF support
- Progress callbacks for async task status updates
- SHA-256 deduplication
- Section title and article number extraction for v2 RAG citation precision
- 100MB file size enforcement
"""

import os
import re
import hashlib
from typing import Callable
from datetime import datetime, timezone

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from db.mongo import documents_col
from services.logger import get_logger

log = get_logger(__name__)

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None  # type: ignore

try:
    import pytesseract
    from pdf2image import convert_from_path  # type: ignore
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None  # type: ignore
    DOCX_AVAILABLE = False

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_SIZE_MB", "100")) * 1024 * 1024  # 100MB


# ── Article number extraction ─────────────────────────────────────────────────
_ARTICLE_PATTERN = re.compile(
    r"(?:Article|Art\.?)\s+(\d+(?:\(\d+\))?(?:\([a-z]\))?)",
    re.IGNORECASE,
)
_SECTION_TITLE_PATTERN = re.compile(
    r"^(Chapter\s+\w+|Section\s+\d+[\.\d]*|CHAPTER\s+\w+|SECTION\s+\d+)",
    re.MULTILINE | re.IGNORECASE,
)


def _extract_article_number(text: str) -> str | None:
    """Extract the first EU AI Act article number referenced in a text chunk."""
    match = _ARTICLE_PATTERN.search(text)
    return f"Article {match.group(1)}" if match else None


def _extract_section_title(text: str) -> str | None:
    """Extract the first section/chapter heading found in a chunk."""
    match = _SECTION_TITLE_PATTERN.search(text)
    return match.group(0).strip() if match else None


def compute_sha256(file_path: str) -> str:
    """Compute SHA-256 hash of a file for deduplication."""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def check_duplicate(sha256_hash: str, org_id: str | None) -> str | None:
    """
    Return existing doc_id if this hash already exists in the org's corpus,
    else return None.
    """
    query: dict = {"sha256_hash": sha256_hash}
    if org_id:
        query["org_id"] = org_id
    existing = documents_col.find_one(query, {"_id": 1})
    return str(existing["_id"]) if existing else None


def ingest_document(
    file_path: str,
    uploaded_by: str = "user",
    org_id: str | None = None,
    is_permanent: bool = False,
    password: str | None = None,
    version: str | None = None,
    source_url: str | None = None,
    supersedes: str | None = None,
    on_progress: Callable[[str], None] | None = None,
) -> tuple[list[dict], str]:
    """
    Ingest a document file into Arkive's knowledge base.

    Args:
        file_path:    Absolute path to the file
        uploaded_by:  User ID of the uploader
        org_id:       Organisation ID for data scoping
        is_permanent: If True, document survives TTL cleanup (knowledge base docs)
        password:     PDF password if encrypted
        version:      Document version string
        source_url:   Origin URL if fetched remotely
        supersedes:   Document ID this version supersedes
        on_progress:  Callback(stage_name) called at each processing stage

    Returns:
        (chunk_data: list[dict], doc_id: str)
    """
    def _progress(stage: str):
        if on_progress:
            on_progress(stage)

    log.info("ingestion_start", extra={"file": os.path.basename(file_path), "org_id": org_id})

    # ── File size check ───────────────────────────────────────────────────────
    file_size = os.path.getsize(file_path)
    if file_size > MAX_UPLOAD_BYTES:
        max_mb = MAX_UPLOAD_BYTES // (1024 * 1024)
        raise ValueError(
            f"File size ({file_size // (1024*1024)}MB) exceeds the {max_mb}MB limit. "
            f"Please split the document or contact support for large document ingestion."
        )

    _progress("extracting_text")

    # ── SHA-256 hash for deduplication ───────────────────────────────────────
    sha256_hash = compute_sha256(file_path)
    log.info("ingestion_hash", extra={"sha256": sha256_hash[:16] + "..."})

    ext = os.path.splitext(file_path)[1].lower()
    pages: list[Document] = []

    if ext == ".pdf":
        # ── PDF text extraction ───────────────────────────────────────────────
        if fitz is None:
            raise RuntimeError("PyMuPDF (fitz) is not installed")

        doc = fitz.open(file_path)
        if doc.needs_pass:
            if password and doc.authenticate(password):
                pass
            else:
                raise ValueError(
                    "PDF is password-protected. Provide the correct password to proceed."
                )

        needs_ocr = False
        for page_num in range(len(doc)):
            text = doc[page_num].get_text("text", sort=True).strip()
            if len(text) > 20:  # Ignore near-empty pages
                pages.append(Document(
                    page_content=text,
                    metadata={"page": page_num + 1}
                ))
            else:
                needs_ocr = True

        # ── OCR fallback for scanned documents ───────────────────────────────
        if needs_ocr and OCR_AVAILABLE:
            log.info("ocr_fallback", extra={"file": os.path.basename(file_path)})
            images = convert_from_path(file_path)
            pages = []
            for i, img in enumerate(images):
                text = pytesseract.image_to_string(img)
                pages.append(Document(page_content=text, metadata={"page": i + 1}))
        elif needs_ocr and not OCR_AVAILABLE:
            log.warning("ocr_unavailable", extra={"note": "Some pages may be empty (scanned PDF)"})

    elif ext == ".docx":
        # ── DOCX extraction ──────────────────────────────────────────────────
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        docx_doc = DocxDocument(file_path)
        full_text = "\n\n".join(
            para.text for para in docx_doc.paragraphs if para.text.strip()
        )
        pages.append(Document(page_content=full_text, metadata={"page": 1}))
        log.info("docx_extracted", extra={"paragraphs": len(docx_doc.paragraphs)})

    elif ext == ".txt":
        # ── Plain text extraction ────────────────────────────────────────────
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()
        pages.append(Document(page_content=raw_text, metadata={"page": 1}))
        log.info("txt_extracted", extra={"chars": len(raw_text)})

    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported formats: PDF, DOCX, TXT."
        )

    log.info("ingestion_pages", extra={"total_pages": len(pages)})
    _progress("chunking")

    # ── Text splitting ────────────────────────────────────────────────────────
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=80,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    chunks = splitter.split_documents(pages)
    log.info("ingestion_chunks", extra={"total_chunks": len(chunks)})

    # ── Save document record to MongoDB ──────────────────────────────────────
    doc_record = {
        "filename": os.path.basename(file_path),
        "file_path": file_path,
        "uploaded_by": uploaded_by,
        "org_id": org_id,
        "uploaded_at": datetime.now(timezone.utc),
        "total_pages": len(pages),
        "total_chunks": len(chunks),
        "sha256_hash": sha256_hash,
        "file_size_bytes": file_size,
        "status": "ingested",
        "is_permanent": is_permanent,
        "version": version,
        "source_url": source_url,
        "supersedes": supersedes,
    }
    result = documents_col.insert_one(doc_record)
    doc_id = str(result.inserted_id)
    log.info("ingestion_saved", extra={"doc_id": doc_id})

    # ── Build chunk data with enhanced metadata ───────────────────────────────
    _progress("embedding")
    chunk_data = []
    for i, chunk in enumerate(chunks):
        text = chunk.page_content
        chunk_data.append({
            "doc_id": doc_id,
            "chunk_index": i,
            "text": text,
            "page": chunk.metadata.get("page", 0),
            "source": os.path.basename(file_path),
            "org_id": org_id,
            "is_permanent": is_permanent,
            # v2 RAG citation fields
            "section_title": _extract_section_title(text),
            "article_number": _extract_article_number(text),
        })

    return chunk_data, doc_id